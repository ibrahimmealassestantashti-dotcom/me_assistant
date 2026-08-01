import streamlit as st
import json
import os
import re
import io
import time
import tempfile
from google import genai
from google.genai import types as genai_types
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# --- إعدادات الصفحة ---
st.set_page_config(
    page_title="مساعد إدارة ومتابعة المشاريع - ME Assistant",
    page_icon="📊",
    layout="wide"
)

CONFIG_FILE = "saved_projects.json"

def _unpack_match_result(stored):
    """فك متوافق مع نتائج المطابقة المخزنة، سواء كانت بالتركيبة القديمة (3 عناصر)
    أو الجديدة (4 عناصر تتضمن سجل التشخيص) - يمنع كسر التطبيق بعد أي تحديث للكود."""
    if len(stored) == 4:
        return stored
    att_v, rep_v, diff_v = stored
    return att_v, rep_v, diff_v, []

def load_saved_projects():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_projects(projects_dict):
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(projects_dict, f, ensure_ascii=False, indent=4)
    except Exception as e:
        st.error(f"خطأ أثناء حفظ البيانات: {e}")

@st.cache_resource
def get_drive_service():
    try:
        if "gcp_service_account" in st.secrets:
            creds_data = dict(st.secrets["gcp_service_account"])
            if "private_key" in creds_data:
                creds_data["private_key"] = creds_data["private_key"].replace("\\n", "\n")

            creds = service_account.Credentials.from_service_account_info(
                creds_data, scopes=["https://www.googleapis.com/auth/drive.readonly"]
            )
            return build("drive", "v3", credentials=creds)
        else:
            st.error("❌ لم يتم العثور على gcp_service_account في Secrets.")
            return None
    except Exception as e:
        st.error(f"❌ خطأ في الاتصال بخدمة Google Drive: {e}")
        return None

def get_folder_contents(service, folder_id):
    query = f"'{folder_id}' in parents and trashed = false"
    folders = []
    files = []
    try:
        results = service.files().list(
            q=query, supportsAllDrives=True, includeItemsFromAllDrives=True,
            fields="files(id, name, mimeType, webViewLink)"
        ).execute()
        
        for item in results.get("files", []):
            if item["mimeType"] == "application/vnd.google-apps.folder":
                folders.append(item)
            else:
                files.append(item)
    except Exception as e:
        st.error(f"خطأ أثناء قراءة المجلد: {e}")
    return folders, files

def download_file_bytes(service, file_item):
    """يحمّل بايتات الملف. يدعم الملفات الثنائية العادية (PDF/صور/Word المرفوعة)
    وأيضاً مستندات Google الأصلية (Google Docs / Google Sheets) عبر التصدير،
    لأن هذه الأخيرة لا يمكن تحميلها مباشرة بواسطة get_media."""
    file_id = file_item.get("id")
    mime_type = file_item.get("mimeType", "")
    try:
        if mime_type.startswith("application/vnd.google-apps"):
            # مستند Google أصلي (تم إنشاؤه داخل درايف مباشرة) - يحتاج تصدير لصيغة PDF
            request = service.files().export_media(fileId=file_id, mimeType="application/pdf")
        else:
            request = service.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        fh.seek(0)
        return fh.read()
    except Exception:
        return None

def is_sub_component(folder_name):
    name = folder_name.lower().strip()
    keywords = [
        "attendance", "attend", "حضور", "كشف", "اسماء", "أسماء",
        "documentation", "doc", "photo", "photos", "image", "images", "صور", "توثيق", "أرشيف",
        "report", "reports", "تقرير", "تقارير", "ملخص"
    ]
    return any(kw in name for kw in keywords)

def parse_session_subfolders(service, session_folder):
    sub_folders, direct_files = get_folder_contents(service, session_folder["id"])
    
    session_data = {
        "session_id": session_folder["id"],
        "session_name": session_folder["name"],
        "attendance": {"folder": None, "files": []},
        "documentation": {"folder": None, "files": []},
        "report": {"folder": None, "files": []},
        "extra_files": direct_files
    }
    
    for sf in sub_folders:
        sf_name_lower = sf["name"].lower().strip()
        _, files = get_folder_contents(service, sf["id"])
        
        if any(k in sf_name_lower for k in ["attendance", "attend", "حضور", "كشف", "اسماء", "أسماء"]):
            session_data["attendance"]["folder"] = sf
            session_data["attendance"]["files"] = files
        elif any(k in sf_name_lower for k in ["documentation", "doc", "photo", "photos", "image", "images", "صور", "توثيق", "أرشيف"]):
            session_data["documentation"]["folder"] = sf
            session_data["documentation"]["files"] = files
        elif any(k in sf_name_lower for k in ["report", "reports", "تقرير", "تقارير", "ملخص"]):
            session_data["report"]["folder"] = sf
            session_data["report"]["files"] = files
            
    return session_data

def fetch_structured_sessions(service, target_folder_id):
    sessions_list = []
    top_folders, _ = get_folder_contents(service, target_folder_id)
    
    for folder in top_folders:
        if is_sub_component(folder["name"]):
            continue
            
        child_folders, _ = get_folder_contents(service, folder["id"])
        has_sub_components = any(is_sub_component(cf["name"]) for cf in child_folders)
        
        if has_sub_components:
            parsed = parse_session_subfolders(service, folder)
            sessions_list.append(parsed)
        else:
            for cf in child_folders:
                if is_sub_component(cf["name"]):
                    continue
                    
                sub_cf_folders, _ = get_folder_contents(service, cf["id"])
                if any(is_sub_component(scf["name"]) for scf in sub_cf_folders):
                    parsed = parse_session_subfolders(service, cf)
                    parsed["session_name"] = f"{folder['name']} / {cf['name']}"
                    sessions_list.append(parsed)
                    
    return sessions_list

def wait_for_file_active(client, g_file, timeout=90):
    """ينتظر حتى تنتهي خوادم Gemini من معالجة الملف المرفوع (PROCESSING -> ACTIVE)
    قبل استخدامه في generate_content، لتفادي فشل المطابقة بسبب ملف لم يجهز بعد."""
    start = time.time()
    while g_file.state.name == "PROCESSING":
        if time.time() - start > timeout:
            raise TimeoutError(f"انتهت مهلة معالجة الملف '{g_file.display_name}' على خوادم Gemini.")
        time.sleep(2)
        g_file = client.files.get(name=g_file.name)
    if g_file.state.name == "FAILED":
        raise ValueError(f"فشلت معالجة الملف '{g_file.display_name}' على خوادم Gemini.")
    return g_file

def extract_text_from_docx_bytes(f_bytes):
    """يستخرج النص من ملف Word محلياً، لأن Gemini File API لا يدعم رفع .docx مباشرة."""
    if DocxDocument is None:
        return None
    doc = DocxDocument(io.BytesIO(f_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # قراءة الجداول أيضاً، لأن بيانات التقارير غالباً تكون داخل جدول
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs)

# --- دالة التحليل الذكي مع التثبيت المزدوج لمفتاح API ومتغير البيئة ---
def extract_session_metrics_with_ai(service, session_info, api_key, model_name):
    if not api_key:
        return ["--/--/--", "0", "0", "0", "0", "0", "0"], ["--/--/--", "0", "0", "0", "0", "0", "0"], ["⚠️ أدخل مفتاح API", "⚠️", "✅", "✅", "✅", "✅", "✅"], ["⚠️ لم يتم إدخال مفتاح API."]

    clean_key = api_key.strip()
    client = genai.Client(api_key=clean_key)

    att_files = session_info.get("attendance", {}).get("files", [])
    rep_files = session_info.get("report", {}).get("files", [])
    
    uploaded_gemini_files = []
    extracted_text_blocks = []
    debug_log = []

    if not att_files and not rep_files:
        debug_log.append("⚠️ لا توجد أي ملفات مكتشفة أصلاً داخل مجلدي الحضور/التقرير لهذه الجلسة.")

    try:
        all_target_files = att_files + rep_files
        for f in all_target_files:
            f_mime = f.get("mimeType", "غير معروف")
            f_bytes = download_file_bytes(service, f)
            if not f_bytes:
                debug_log.append(f"❌ فشل تحميل الملف '{f['name']}' (النوع: {f_mime}) من Drive.")
                continue
            else:
                debug_log.append(f"✅ تم تحميل الملف '{f['name']}' بنجاح ({len(f_bytes)} bytes، النوع: {f_mime}).")

            fname_lower = f["name"].lower()
            is_google_native = f_mime.startswith("application/vnd.google-apps")

            # نحدد "النوع الفعلي" اعتماداً على mimeType الحقيقي القادم من Drive أولاً
            # (بعض الملفات المرفوعة لا تحمل امتداداً في اسمها، مثل "Report" بدون .pdf)
            # ونستخدم امتداد الاسم فقط كخيار احتياطي إن كان mimeType غامضاً
            is_pdf = is_google_native or f_mime == "application/pdf" or fname_lower.endswith(".pdf")
            is_png = f_mime == "image/png" or fname_lower.endswith(".png")
            is_jpg = f_mime in ("image/jpeg", "image/jpg") or fname_lower.endswith((".jpg", ".jpeg"))
            is_docx = (not is_pdf and not is_png and not is_jpg) and (
                f_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                or fname_lower.endswith(".docx")
            )
            is_old_doc = (not is_pdf and not is_png and not is_jpg and not is_docx) and (
                f_mime == "application/msword" or fname_lower.endswith(".doc")
            )

            # مستندات Google الأصلية (Docs/Sheets) تم تصديرها أعلاه كـ PDF فعلياً بغض النظر عن اسمها
            if is_pdf:
                mime_type = "application/pdf"
            # ملفات Word: Gemini File API لا يدعم رفع .docx/.doc مباشرة (يرجع خطأ Unsupported MIME type)
            # لذلك نستخرج النص محلياً ونرسله كنص ضمن البرومبت بدل رفعه كملف
            elif is_docx:
                try:
                    doc_text = extract_text_from_docx_bytes(f_bytes)
                    if doc_text is None:
                        extracted_text_blocks.append(
                            f"⚠️ لم تتمكن من قراءة ملف Word ({f['name']}) لأن مكتبة python-docx غير مثبتة في البيئة."
                        )
                    else:
                        extracted_text_blocks.append(f"محتوى ملف Word ({f['name']}):\n{doc_text}")
                except Exception as docx_err:
                    extracted_text_blocks.append(f"⚠️ تعذرت قراءة ملف Word ({f['name']}): {docx_err}")
                continue
            elif is_old_doc:
                extracted_text_blocks.append(
                    f"⚠️ الملف ({f['name']}) بصيغة .doc القديمة غير مدعومة تلقائياً، يفضّل حفظه كـ .docx أو PDF."
                )
                continue
            elif is_png:
                mime_type = "image/png"
            elif is_jpg:
                mime_type = "image/jpeg"
            else:
                # نوع غير مدعوم من Gemini File API لعرض المستندات - نتجاهله بدل ما يفشل التحليل كامل
                extracted_text_blocks.append(f"⚠️ نوع الملف ({f['name']}) غير مدعوم للتحليل التلقائي، تم تجاهله.")
                debug_log.append(f"⏭️ تم تجاهل '{f['name']}' لأن نوعه ({f_mime}) غير مدعوم.")
                continue

            suffix = ".pdf" if (is_google_native or is_pdf) else os.path.splitext(f["name"])[1] or (".png" if is_png else ".jpg")
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(f_bytes)
                tmp_path = tmp.name
            
            # رفع الملف ثم الانتظار حتى تكتمل معالجته على خوادم Gemini (PROCESSING -> ACTIVE)
            g_file = client.files.upload(
                file=tmp_path,
                config=genai_types.UploadFileConfig(mime_type=mime_type, display_name=f["name"]),
            )
            g_file = wait_for_file_active(client, g_file)
            uploaded_gemini_files.append(g_file)
            debug_log.append(f"📤 تم رفع '{f['name']}' إلى Gemini وأصبح جاهزاً للتحليل (state: {g_file.state.name}).")
            try:
                os.remove(tmp_path)
            except:
                pass

        prompt = f"""
        أنت مدقق بيانات مشاريع محترف. قم بقراءة وفحص مستندات الحضور والتقارير المرفقة لهذه الجلسة ("{session_info.get('session_name')}") بدقة متناهية.
        ملاحظات هامة جداً:
        1. ورقة الحضور ملف PDF ممسوح ضوئياً والأعداد والتواريخ مكتوبة بخط اليد.
        2. التقرير ملف وورد (Word) أو PDF والبيانات المطلوبة موجودة في النصف الأول من الصفحة الأولى.
        3. عدم وجود رقم أو خانة فارغة يعني تماماً أن القيمة هي صفر (0).
        
        استخرج البيانات التالية بدقة على شكل JSON صارم:
        - attendance_data: [التاريخ، العدد الإجمالي، الرجال (Men)، النساء (Women)، الأولاد/أطفال ذكور (Boys)، الفتيات (Girls)، ذوي الاحتياجات (PWD)]
        - report_data: [التاريخ، العدد الإجمالي، الرجال (Men)، النساء (Women)، الأولاد/أطفال ذكور (Boys)، الفتيات (Girls)، ذوي الاحتياجات (PWD)]
        - differences: [تقييم لكل بند: اكتب "مطابقة" إذا كانت متطابقة تماماً، أو اكتب وصف الفارق إذا وُجد اختلاف]

        أجب بصيغة JSON صارمة فقط وبدون أي نصوص إضافية، بحيث تكون القوائم تحتوي على 7 عناصر تماماً، وجميع القيم الرقمية أو التواريخ مضبوطة وصحيحة. وإذا لم توجد قيمة ضع "0".
        """

        if extracted_text_blocks:
            prompt += "\n\n--- نصوص مستخرجة مسبقاً من ملفات Word المرفقة (اعتمد عليها بدل الصورة) ---\n"
            prompt += "\n\n".join(extracted_text_blocks)

        debug_log.append(f"📊 الملخص: {len(uploaded_gemini_files)} ملف تم رفعه للذكاء الاصطناعي، {len(extracted_text_blocks)} كتلة نص مستخرجة من Word/تحذيرات.")

        response = client.models.generate_content(model=model_name, contents=[prompt] + uploaded_gemini_files)
        text_res = response.text
        
        for g_file in uploaded_gemini_files:
            try:
                client.files.delete(name=g_file.name)
            except:
                pass

        match = re.search(r'\{.*\}', text_res, re.DOTALL)
        if match:
            data = json.loads(match.group(0))
            return (
                data.get("attendance_data", ["--", "0", "0", "0", "0", "0", "0"]),
                data.get("report_data", ["--", "0", "0", "0", "0", "0", "0"]),
                data.get("differences", ["✅", "✅", "✅", "✅", "✅", "✅", "✅"]),
                debug_log,
            )
        else:
            debug_log.append("⚠️ رد الذكاء الاصطناعي لم يحتوِ على JSON صالح.")
            debug_log.append(f"نص الرد الخام (أول 500 حرف): {text_res[:500]}")

    except Exception as e:
        st.error(f"خطأ أثناء معالجة وتحليل الملفات: {e}")
        debug_log.append(f"❌ استثناء: {e}")
        for g_file in uploaded_gemini_files:
            try:
                client.files.delete(name=g_file.name)
            except:
                pass
        
    return ["--/--/--", "0", "0", "0", "0", "0", "0"], ["--/--/--", "0", "0", "0", "0", "0", "0"], ["⚠️ تعذر التحليل", "⚠️ تعذر التحليل", "✅ مطابق", "✅ مطابق", "✅ مطابق", "✅ مطابق", "✅ مطابق"], debug_log

# --- الواجهة الرئيسية ---
st.title("📊 نظام إدارة ومتابعة المشاريع الذكي (ME Assistant)")
st.markdown("---")

if "projects" not in st.session_state:
    st.session_state.projects = load_saved_projects()

# --- القائمة الجانبية (Sidebar) ---
with st.sidebar:
    st.header("🔑 إعدادات الذكاء الاصطناعي")
    user_gemini_key = st.text_input("مفتاح Gemini API", type="password", value=st.secrets.get("GEMINI_API_KEY", ""))
    
    free_models_options = [
        "gemini-3.6-flash",
        "gemini-3.5-flash",
        "gemini-3.5-flash-lite",
        "gemini-3.1-pro-preview",
        "gemini-3.1-flash-lite"
    ]
    
    selected_ai_model = st.selectbox("اختر نموذج الذكاء الاصطناعي:", free_models_options, index=0)
    final_model_to_use = selected_ai_model

    if st.button("🧪 فحص الأداة والمفتاح"):
        if not user_gemini_key:
            st.warning("⚠️ يرجى إدخال مفتاح API أولاً.")
        else:
            with st.spinner("جاري فحص الاتصال..."):
                try:
                    c_key = user_gemini_key.strip()
                    test_client = genai.Client(api_key=c_key)
                    test_res = test_client.models.generate_content(model=final_model_to_use, contents="مرحبا")
                    if test_res.text:
                        st.success(f"✅ الاتصال ناجح والنموذج ({final_model_to_use}) يعمل بكفاءة.")
                except Exception as e:
                    st.error(f"❌ فشل الفحص: {e}")

    st.markdown("---")
    st.header("⚙️ إدارة المشاريع المحفوظة")
    new_project_name = st.text_input("اسم المشروع")
    new_project_id = st.text_input("معرف المجلد (Folder ID)")
    
    if st.button("➕ إضافة وحفظ المشروع"):
        if new_project_name and new_project_id:
            st.session_state.projects[new_project_name] = new_project_id.strip()
            save_projects(st.session_state.projects)
            st.success(f"تم حفظ مشروع '{new_project_name}' بنجاح!")
            st.rerun()
        else:
            st.error("يرجى إدخال البيانات كاملة.")
            
    if st.session_state.projects:
        st.markdown("---")
        st.subheader("📋 المشاريع المسجلة:")
        for p_name in list(st.session_state.projects.keys()):
            col_p1, col_p2 = st.columns([3, 1])
            col_p1.text(p_name)
            if col_p2.button("🗑️", key=f"del_{p_name}"):
                del st.session_state.projects[p_name]
                save_projects(st.session_state.projects)
                st.rerun()

service = get_drive_service()

if not st.session_state.projects:
    st.info("👈 قم بإضافة أول مشروع من الشريط الجانبي وسيبقى محفوظاً دائماً.")
else:
    project_names = list(st.session_state.projects.keys())
    selected_tab = st.tabs(project_names)

    for idx, p_name in enumerate(project_names):
        with selected_tab[idx]:
            root_id = st.session_state.projects[p_name]
            
            path_key = f"path_{p_name}"
            if path_key not in st.session_state:
                st.session_state[path_key] = [{"id": root_id, "name": p_name}]
                
            current_trail = st.session_state[path_key]
            current_folder = current_trail[-1]
            
            st.subheader(f"📁 متصفح مجلدات مشروع: {p_name}")
            trail_str = " ➡️ ".join([node["name"] for node in current_trail])
            st.info(f"📍 **المسار الحالي:** {trail_str}")
            
            col_nav1, col_nav2 = st.columns([1, 4])
            with col_nav1:
                if len(current_trail) > 1:
                    if st.button("⬅️ العودة للمجلد السابق", key=f"back_{p_name}"):
                        st.session_state[path_key].pop()
                        st.rerun()
                        
            if service:
                sub_folders, direct_files = get_folder_contents(service, current_folder["id"])
                
                if sub_folders:
                    folder_options = {sf["name"]: sf["id"] for sf in sub_folders}
                    selected_sub_name = st.selectbox(
                        "اختر المجلد/النشاط للانتقال إليه:",
                        ["-- اختر المجلد --"] + list(folder_options.keys()),
                        key=f"select_folder_{p_name}"
                    )
                    
                    if selected_sub_name != "-- اختر المجلد --":
                        next_id = folder_options[selected_sub_name]
                        st.session_state[path_key].append({"id": next_id, "name": selected_sub_name})
                        st.rerun()

                st.markdown("---")
                
                btn_fetch = st.button(
                    f"⚡ جلب وتحليل الجلسات المندمجة تحت ({current_folder['name']})", 
                    key=f"fetch_btn_{p_name}",
                    type="primary"
                )

                if btn_fetch:
                    with st.spinner("جاري تحديد مجلدات الجلسات وفحص المجلدات الفرعية..."):
                        sessions = fetch_structured_sessions(service, current_folder["id"])
                        st.session_state[f"data_{p_name}"] = sessions

                sessions_data = st.session_state.get(f"data_{p_name}", None)
                
                if sessions_data is not None:
                    if not sessions_data:
                        st.warning("لم يتم العثور على جلسات مكتملة الهيكلية داخل هذا المجلد.")
                    else:
                        st.success(f"تم اكتشاف {len(sessions_data)} جلسة/جلسات بالهيكلية الصحيحة!")
                        st.markdown("---")
                        
                        st.subheader("🛠️ أدوات التحليل والمطابقة الذكية")
                        b1, b2, b3, b4 = st.columns(4)
                        
                        if b1.button("1️⃣ فحص المرفقات للجلسات", key=f"b1_{p_name}"):
                            st.session_state[f"view_{p_name}"] = "attachments"
                        if b2.button("2️⃣ مطابقة الحضور والتقرير", key=f"b2_{p_name}"):
                            st.session_state[f"view_{p_name}"] = "matching"
                        if b3.button("3️⃣ إحصائية الفئات والحضور", key=f"b3_{p_name}"):
                            st.session_state[f"view_{p_name}"] = "stats"
                        if b4.button("4️⃣ المساعد الذكي (AI Chat)", key=f"b4_{p_name}"):
                            st.session_state[f"view_{p_name}"] = "chat"

                        current_view = st.session_state.get(f"view_{p_name}", None)

                        if current_view == "attachments":
                            st.markdown("#### 📋 نتيجة فحص مجلدات الجلسات الفرعية الثلاثة:")
                            for s_idx, sess in enumerate(sessions_data):
                                att_files = sess.get("attendance", {}).get("files", [])
                                doc_files = sess.get("documentation", {}).get("files", [])
                                rep_files = sess.get("report", {}).get("files", [])
                                
                                is_complete = len(att_files) > 0 and len(doc_files) > 0 and len(rep_files) > 0
                                status_tag = "✅ مكتملة" if is_complete else "⚠️ ناقصة"
                                sess_title = sess.get("session_name", "جلسة بدون عنوان")
                                
                                with st.expander(f"📌 **{sess_title}** — الحالة: {status_tag}"):
                                    col_f1, col_f2, col_f3 = st.columns(3)
                                    with col_f1:
                                        st.write("**📄 ورقة الحضور:**")
                                        for f in att_files: st.caption(f"• {f['name']}")
                                    with col_f2:
                                        st.write("**🖼️ صور التوثيق:**")
                                        for f in doc_files: st.caption(f"• {f['name']}")
                                    with col_f3:
                                        st.write("**📑 التقرير:**")
                                        for f in rep_files: st.caption(f"• {f['name']}")

                        elif current_view == "matching":
                            st.markdown("#### ⚖️ مطابقة أوراق الحضور مع التقارير بقراءة الملفات الحقيقية:")
                            for s_idx, sess in enumerate(sessions_data):
                                sess_title = sess.get("session_name", "جلسة")
                                
                                col_m1, col_m2 = st.columns([4, 1])
                                col_m1.markdown(f"##### 📌 جلسة: {sess_title}")
                                
                                match_btn_key = f"match_exec_{p_name}_{s_idx}"
                                result_key = f"match_res_{p_name}_{s_idx}"
                                
                                if col_m2.button("🔍 مطابقة الجلسة", key=match_btn_key):
                                    with st.spinner("جاري قراءة وتحليل المستندات الفعليّة بدقة..."):
                                        att_v, rep_v, diff_v, debug_v = extract_session_metrics_with_ai(service, sess, user_gemini_key, final_model_to_use)
                                        st.session_state[result_key] = (att_v, rep_v, diff_v, debug_v)

                                if result_key in st.session_state:
                                    att_vals, rep_vals, diff_vals, debug_vals = _unpack_match_result(st.session_state[result_key])

                                    with st.expander("🔧 تفاصيل المعالجة (تشخيص)", expanded=False):
                                        if debug_vals:
                                            for line in debug_vals:
                                                st.caption(line)
                                        else:
                                            st.caption("لا توجد بيانات تشخيص.")
                                    
                                    def safe_pad(lst, target_len=7):
                                        return lst + ["0"] * (target_len - len(lst)) if len(lst) < target_len else lst[:target_len]

                                    att_vals = safe_pad(att_vals)
                                    rep_vals = safe_pad(rep_vals)
                                    diff_vals = safe_pad(diff_vals)

                                    comparison_data = {
                                        "البند": ["تاريخ الجلسة", "العدد الإجمالي", "الرجال (Men)", "النساء (Women)", "الأولاد (Boys)", "الفتيات (Girls)", "ذوي الاحتياجات (PWD)"],
                                        "ورقة الحضور": att_vals,
                                        "التقرير": rep_vals,
                                        "النتيجة / الفروقات": diff_vals
                                    }
                                    st.table(comparison_data)
                                else:
                                    st.info("اضغط على زر (مطابقة الجلسة) أعلاه لبدء قراءة الملفات وتحليلها.")
                                st.markdown("---")

                        elif current_view == "stats":
                            st.markdown("#### 📊 الإحصائية التجميعية للمستفيدين عبر الجلسات المطابقة:")
                            tot_men, tot_women, tot_boys, tot_girls, tot_pwd = 0, 0, 0, 0, 0
                            
                            for s_idx, s in enumerate(sessions_data):
                                res_key = f"match_res_{p_name}_{s_idx}"
                                if res_key in st.session_state:
                                    att_v, _, _, _ = _unpack_match_result(st.session_state[res_key])
                                    try:
                                        tot_men += int(att_v[2]) if att_v[2].isdigit() else 0
                                        tot_women += int(att_v[3]) if att_v[3].isdigit() else 0
                                        tot_boys += int(att_v[4]) if att_v[4].isdigit() else 0
                                        tot_girls += int(att_v[5]) if att_v[5].isdigit() else 0
                                        tot_pwd += int(att_v[6]) if att_v[6].isdigit() else 0
                                    except Exception:
                                        pass
                                
                            col_stat1, col_stat2, col_stat3, col_stat4, col_stat5 = st.columns(5)
                            col_stat1.metric("👨 رجال", str(tot_men))
                            col_stat2.metric("👩 نساء", str(tot_women))
                            col_stat3.metric("👧 فتيات إناث", str(tot_girls))
                            col_stat4.metric("👶 أطفال ذكور", str(tot_boys))
                            col_stat5.metric("♿ ذوي الاحتياجات", str(tot_pwd))

                        elif current_view == "chat":
                            st.markdown("---")
                            st.subheader("💬 دردشة المساعد الذكي لمراجعة الجلسات")

                            chat_history_key = f"messages_{p_name}"
                            if chat_history_key not in st.session_state:
                                st.session_state[chat_history_key] = []

                            for message in st.session_state[chat_history_key]:
                                with st.chat_message(message["role"]):
                                    st.write(message["content"])

                            if prompt_text := st.chat_input("وجه سؤالك للذكاء الاصطناعي حول بيانات وقراءات الجلسات..."):
                                st.session_state[chat_history_key].append({"role": "user", "content": prompt_text})
                                with st.chat_message("user"):
                                    st.write(prompt_text)

                                context_text = f"مشروع: {p_name} | المجلد: {current_folder['name']}\n"
                                context_text += f"عدد الجلسات الإجمالي: {len(sessions_data)}\n\n"
                                
                                for s_idx, s in enumerate(sessions_data):
                                    s_title = s.get('session_name', '')
                                    res_key = f"match_res_{p_name}_{s_idx}"
                                    if res_key in st.session_state:
                                        att_v, rep_v, _, _ = _unpack_match_result(st.session_state[res_key])
                                        context_text += f"- {s_title}:\n"
                                        try:
                                            context_text += f"  * التواريخ: حضور ({att_v[0]})، تقرير ({rep_v[0]})\n"
                                            context_text += f"  * المجموع: حضور ({att_v[1]})، تقرير ({rep_v[1]})\n"
                                        except IndexError:
                                            pass

                                with st.chat_message("assistant"):
                                    with st.spinner("جاري معالجة السؤال..."):
                                        try:
                                            c_key = user_gemini_key.strip()
                                            chat_client = genai.Client(api_key=c_key)
                                            ai_prompt = f"أنت مساعد ذكي لإدارة المشاريع.\nاجب بلغة عربية دقيقة بناءً على البيانات التالية:\n{context_text}\nسؤال المستخدم: {prompt_text}"
                                            response = chat_client.models.generate_content(model=final_model_to_use, contents=ai_prompt)
                                            ai_response = response.text
                                            st.write(ai_response)
                                            st.session_state[chat_history_key].append({"role": "assistant", "content": ai_response})
                                        except Exception as e:
                                            st.error(f"❌ فشل الاتصال: {e}")
            else:
                st.warning("⚠️ لا يوجد اتصال بخدمة Google Drive.")

st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: #888888; padding: 10px;'>"
    "تم تصميم وتطوير البرنامج بواسطة <b>إبراهيم الجاسم</b> © 2026"
    "</div>", 
    unsafe_allow_html=True
)
